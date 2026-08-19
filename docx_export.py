# Copyright (C) 2026 Ori Mosenzon and Claude (Anthropic AI)
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
# See the LICENSE file for details.

"""Build a Word (.docx) document from a cached song: lyrics plus optional translation."""

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Credit labels, mirroring CREDITS_LABELS in templates/index.html.
CREDITS_LABELS = {
    "he": {"lyrics": "מילים", "music": "לחן", "both": "מילים ולחן", "arranger": "עיבוד", "performer": "מבצע"},
    "en": {"lyrics": "Lyrics", "music": "Music", "both": "Music & Lyrics", "arranger": "Arr.", "performer": "By"},
    "ar": {"lyrics": "كلمات", "music": "ألحان", "both": "كلمات وألحان", "arranger": "توزيع", "performer": "أداء"},
    "ru": {"lyrics": "Слова", "music": "Музыка", "both": "Слова и музыка", "arranger": "Аранж.", "performer": "Исп."},
    "fr": {"lyrics": "Paroles", "music": "Musique", "both": "Paroles et Musique", "arranger": "Arr.", "performer": "Par"},
    "es": {"lyrics": "Letra", "music": "Música", "both": "Letra y Música", "arranger": "Arr.", "performer": "Por"},
    "it": {"lyrics": "Testo", "music": "Musica", "both": "Testo e Musica", "arranger": "Arr.", "performer": "Perf."},
    "de": {"lyrics": "Text", "music": "Musik", "both": "Text und Musik", "arranger": "Arr.", "performer": "Von"},
    "pt": {"lyrics": "Letra", "music": "Música", "both": "Letra e Música", "arranger": "Arr.", "performer": "Por"},
    "ja": {"lyrics": "作詞", "music": "作曲", "both": "作詞・作曲", "arranger": "編曲", "performer": "歌"},
    "ko": {"lyrics": "작사", "music": "작곡", "both": "작사·작곡", "arranger": "편곡", "performer": "노래"},
    "zh": {"lyrics": "词曲", "music": "作曲", "both": "词曲", "arranger": "编曲", "performer": "演唱"},
}

_RTL_CHARS = re.compile(r"[֐-׿؀-ۿ܀-ݏיִ-﷿ﹰ-﻿]")

# Order of child elements inside <w:pPr> per the OOXML schema; w:bidi must precede these.
_PPR_AFTER_BIDI = ("w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing",
                   "w:mirrorIndents", "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
                   "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr",
                   "w:pPrChange")
# Same for <w:rPr>; w:rtl must precede these.
_RPR_AFTER_RTL = ("w:cs", "w:em", "w:lang", "w:eastAsianLayout", "w:specVanish", "w:oMath")


def is_rtl(text: str) -> bool:
    """True when the string contains Hebrew/Arabic letters, so it should be laid out right-to-left."""
    return bool(_RTL_CHARS.search(text or ""))


def _set_para_rtl(paragraph, rtl: bool):
    """Set the paragraph's base direction, so Word aligns and orders it correctly.

    The paragraph direction decides which edge the line hugs; each run is tagged by its own
    script, so an English translation under a Hebrew line still stays on the Hebrew margin.
    """
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1" if rtl else "0")
    pPr.insert_element_before(bidi, *_PPR_AFTER_BIDI)
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        el = OxmlElement("w:rtl")
        el.set(qn("w:val"), "1" if is_rtl(run.text) else "0")
        rPr.insert_element_before(el, *_RPR_AFTER_RTL)


def _set_font(run, name: str):
    """Set the font for Latin, East-Asian and complex-script (Hebrew/Arabic) text alike."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def _para(doc, text, *, size=12, bold=False, italic=False, color=None, align=None,
          space_before=0, space_after=4, font="Arial", rtl=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    _set_para_rtl(p, is_rtl(text) if rtl is None else rtl)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def credits_line(data: dict) -> str:
    """Compose the 'By X · Music Y · Lyrics Z' line, mirroring renderCredits() in the browser."""
    labels = CREDITS_LABELS.get(data.get("lang") or "en", CREDITS_LABELS["en"])
    lyricist = data.get("lyricist")
    composer = data.get("composer")
    parts = []
    if data.get("performer"):
        parts.append(f"{labels['performer']}: {data['performer']}")
    if lyricist and composer and lyricist == composer:
        parts.append(f"{labels['both']}: {lyricist}")
    else:
        if composer:
            parts.append(f"{labels['music']}: {composer}")
        if lyricist:
            parts.append(f"{labels['lyrics']}: {lyricist}")
    if data.get("arranger"):
        parts.append(f"{labels['arranger']}: {data['arranger']}")
    return " · ".join(parts)


def safe_filename(title: str, lang: str = "") -> str:
    """Turn a song title into a filename Windows, macOS and Linux all accept."""
    name = re.sub(r'[\\/:*?"<>|\r\n\t]', " ", title or "lyrics")
    name = re.sub(r"\s+", " ", name).strip(" .")
    name = name[:80].strip() or "lyrics"
    if lang:
        name += f" ({lang})"
    return name + ".docx"


def build_docx(data: dict, translations: list = None, target_lang: str = "") -> io.BytesIO:
    """Render a song into a .docx in memory.

    `translations` is the per-segment translated lines (same length as segments), or None.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    title = data.get("title") or "Lyrics"
    _para(doc, title, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    credits = credits_line(data)
    if credits:
        _para(doc, credits, size=10, color=(0x66, 0x66, 0x66),
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    if target_lang:
        lang_note = f"{data.get('lang', '?')} → {target_lang}"
        _para(doc, lang_note, size=9, italic=True, color=(0x99, 0x99, 0x99),
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, rtl=False)

    if data.get("url"):
        _para(doc, data["url"], size=9, color=(0x0D, 0x94, 0x88),
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14, rtl=False)

    segments = data.get("segments") or []
    for i, seg in enumerate(segments):
        text = (seg.get("text") or "").strip()
        translated = (translations[i] or "").strip() if translations and i < len(translations) else ""
        if not text and not translated:
            _para(doc, "", size=6, space_after=0)
            continue
        # Both lines of a pair take the original's direction, so they stay on the same margin.
        line_rtl = is_rtl(text) if text else is_rtl(translated)
        _para(doc, text, size=12, space_before=6 if translated else 2, space_after=0, rtl=line_rtl)
        if translated and translated != text:
            _para(doc, translated, size=10.5, italic=True, color=(0x77, 0x77, 0x77),
                  space_after=2, rtl=line_rtl)

    _para(doc, "", size=8, space_after=0)
    _para(doc, "Letras · https://github.com/orimosenzon/letras", size=8,
          color=(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER, rtl=False)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
